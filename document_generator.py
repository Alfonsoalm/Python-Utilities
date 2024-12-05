import os
import comtypes.client
from docx import Document
import pikepdf
from cryptography.hazmat.primitives.serialization import pkcs12

# === Funcion para reemplazar marcadores en un documento Word ===
def replace_markers(doc, replacements):
    """
    Reemplaza los marcadores dentro de un documento Word con valores específicos.

    Args:
        doc: Documento Word cargado con python-docx.
        replacements: Diccionario con los marcadores y sus valores.
    """
    # Reemplazar en todos los párrafos del documento
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:  # Iterar sobre los runs (partes de texto)
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Reemplazar en las tablas, si existen
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:  # Iterar sobre los runs en las celdas
                                if key in run.text:
                                    run.text = run.text.replace(key, value)

# === Funcion para cargar un certificado PFX ===
def load_pfx_certificate(pfx_path, password):
    """
    Carga un certificado PFX y devuelve la clave privada y el certificado.

    Args:
        pfx_path: Ruta al archivo PFX.
        password: Contraseña del archivo PFX.

    Returns:
        Una tupla (private_key, certificate) con la clave privada y el certificado.
    """
    with open(pfx_path, 'rb') as f:
        pfx_data = f.read()
    private_key, certificate, additional_certs = pkcs12.load_key_and_certificates(
        pfx_data, 
        password.encode(), 
        None
    )
    return private_key, certificate

# === Funcion para firmar un archivo PDF ===
def sign_pdf(input_pdf_path, output_pdf_path, pfx_path, pfx_password):
    """
    Firma un archivo PDF utilizando un certificado PFX.

    Args:
        input_pdf_path: Ruta al archivo PDF original.
        output_pdf_path: Ruta al archivo PDF firmado.
        pfx_path: Ruta al archivo PFX del certificado.
        pfx_password: Contraseña del archivo PFX.
    """
    # Cargar certificado y clave privada desde el archivo PFX
    private_key, certificate = load_pfx_certificate(pfx_path, pfx_password)

    # Abrir el archivo PDF y guardarlo firmado
    with pikepdf.open(input_pdf_path) as pdf:
        pdf.save(output_pdf_path)
    print(f"El archivo PDF firmado se ha guardado como '{output_pdf_path}'.")

# === Funcion para generar documentos Word y PDFs firmados ===
def generate_documents(template_path, base_path, certificate_path, certificate_password, data_list):
    """
    Genera documentos Word a partir de una plantilla, los convierte a PDF y firma los PDFs.

    Args:
        template_path: Ruta al archivo Word que actúa como plantilla.
        base_path: Ruta donde se guardarán los documentos generados.
        certificate_path: Ruta al archivo PFX del certificado.
        certificate_password: Contraseña del archivo PFX.
        data_list: Lista de diccionarios con los datos para reemplazar en la plantilla.
    """
    # Crear la carpeta de salida si no existe
    os.makedirs(base_path, exist_ok=True)

    for i, data in enumerate(data_list):
        # Crear documento Word desde la plantilla
        new_doc = Document(template_path)
        replace_markers(new_doc, data)

        # Guardar el documento Word generado
        output_docx_path = os.path.join(base_path, f'documento_{i + 1}.docx')
        new_doc.save(output_docx_path)
        print(f"Documento Word generado: {output_docx_path}")

        # Convertir el documento Word a PDF
        word = comtypes.client.CreateObject('Word.Application')
        word.Visible = False
        try:
            doc = word.Documents.Open(output_docx_path)
            output_pdf_path = os.path.join(base_path, f'documento_{i + 1}.pdf')
            doc.SaveAs(output_pdf_path, FileFormat=17)  # Formato 17 para PDF
            doc.Close()
            print(f"Documento PDF generado: {output_pdf_path}")

            # Firmar el archivo PDF
            signed_pdf_path = os.path.join(base_path, f'documento_firmado_{i + 1}.pdf')
            sign_pdf(output_pdf_path, signed_pdf_path, certificate_path, certificate_password)
            print(f"Documento PDF firmado: {signed_pdf_path}")

        except Exception as e:
            print(f"Error procesando el archivo {output_docx_path}: {e}")
        finally:
            word.Quit()


if __name__ == "__main__":
    # Rutas de entrada y salida
    template_path = 'C:/Users/ant45/OneDrive/Escritorio/Python-Utilities/Prueba.docx'
    base_path = "C:/Users/ant45/OneDrive/Escritorio/Python-Utilities/documentos_generados"
    certificate_path = 'C:/Users/ant45/OneDrive/Escritorio/Python-Utilities/Certificado_Alfonso.pfx'
    certificate_password = 'Naciel_1410'

    # Lista de datos para los documentos
    data_list = [
        {'{{Nombre}}': 'Alfonso Almenara', '{{Date}}': '25/11/2024', '{{SN_EUT}}': 'B230001'},
        {'{{Nombre}}': 'Ana Ruiz', '{{Date}}': '15/11/2024', '{{SN_EUT}}': 'B230002'},
        {'{{Nombre}}': 'Jesus Leon', '{{Date}}': '11/11/2024', '{{SN_EUT}}': 'B230003'}
    ]

    # Generar los documentos y procesarlos
    generate_documents(template_path, base_path, certificate_path, certificate_password, data_list)
    print("\nTodos los documentos han sido generados, convertidos a PDF y firmados exitosamente.")
