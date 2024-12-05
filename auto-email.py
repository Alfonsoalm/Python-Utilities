import imaplib
import email
from email.header import decode_header
from credentials import username, password  # Importar credenciales
from fpdf import FPDF
import docx
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from credentials import username, password  # Credenciales de correo

# === Funcion para enviar emails ===
def send_email(to_email, subject, message, cc_email=None, bcc_email=None):
    """
    Envía un correo a un único destinatario con CC y BCC.
    """
    try:
        # Crear el mensaje de correo
        msg = MIMEMultipart()
        msg["From"] = username
        msg["To"] = to_email
        msg["Subject"] = subject

        if cc_email:
            msg["Cc"] = cc_email

        # Agregar el contenido del mensaje
        msg.attach(MIMEText(message, "plain"))

        # Crear la lista de destinatarios
        recipients = [to_email]
        if cc_email:
            recipients.append(cc_email)
        if bcc_email:
            recipients.append(bcc_email)

        # Enviar el correo
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(username, password)
            server.sendmail(username, recipients, msg.as_string())

        print("Correo enviado exitosamente.")
    except Exception as e:
        print(f"Error al enviar el correo: {e}")

# === Funcion para enviar emails a varias personas ===
def send_bulk_emails(to_emails, subject, message):
    """
    Envía un correo a múltiples destinatarios.
    """
    try:
        # Crear el mensaje de correo base
        msg = MIMEMultipart()
        msg["From"] = username
        msg["Subject"] = subject

        # Agregar el contenido del mensaje
        msg.attach(MIMEText(message, "plain"))

        # Enviar el correo a cada destinatario
        with smtplib.SMTP("smtp.gmail.com", 587) as server:
            server.starttls()
            server.login(username, password)
            for email in to_emails:
                msg["To"] = email
                server.sendmail(username, email, msg.as_string())

        print("Correos enviados exitosamente.")
    except Exception as e:
        print(f"Error al enviar correos: {e}")

# === Funcion exportar los emails a un archivo con formato especificado ===
def export_messages(content_list, format, save_path):
    """Guarda todos los mensajes en un solo archivo en el formato especificado."""
    if format == "txt":
        with open(f"{save_path}/emails.txt", "w", encoding="utf-8") as file:
            for content in content_list:
                file.write(content + "\n\n" + ("-" * 80) + "\n\n")
    elif format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for content in content_list:
            # Dividir encabezado y cuerpo
            header, body = content.split("----", 1)
            pdf.set_font("Arial", size=14, style="B")  # Título en negrita y más grande
            for line in header.splitlines():
                pdf.cell(0, 10, line, ln=True)
            pdf.set_font("Arial", size=12)  # Texto del cuerpo más pequeño
            pdf.cell(0, 10, "", ln=True)  # Espaciado
            for line in body.splitlines():
                pdf.cell(0, 10, line, ln=True)
            pdf.cell(0, 10, "-" * 80, ln=True)  # Separador entre mensajes
        pdf.output(f"{save_path}/emails.pdf")
    elif format == "word":
        doc = docx.Document()
        for content in content_list:
            # Dividir encabezado y cuerpo
            header, body = content.split("----", 1)
            # Agregar encabezado con estilo
            title = doc.add_paragraph(header)
            title.style = doc.styles["Heading 1"]  # Usar un estilo grande y en negrita
            # Agregar cuerpo
            doc.add_paragraph(body)
            # Agregar separador
            doc.add_paragraph("-" * 80)
        doc.save(f"{save_path}/emails.docx")

# === Funcion para leer emails ===
def read_emails(username, password, number_emails, format="txt", filter_email=None, save_path="."):
    try:
        # Conectarse al servidor de Gmail
        imap = imaplib.IMAP4_SSL("imap.gmail.com")

        # Iniciar sesión en la cuenta
        imap.login(username, password)

        # Seleccionar la bandeja de entrada
        imap.select("inbox")

        # Buscar correos electrónicos (todos los correos)
        status, messages = imap.search(None, "ALL")

        # Convertir los resultados en una lista de IDs de mensajes
        mail_ids = messages[0].split()

        print(f"Se encontraron {len(mail_ids)} correos electrónicos.")

        content_list = []  # Lista para almacenar los contenidos de los correos

        # Leer los últimos correos electrónicos
        for i in mail_ids[-number_emails:]:
            # Fetch del correo electrónico por su ID
            res, msg = imap.fetch(i, "(RFC822)")

            for response in msg:
                if isinstance(response, tuple):
                    # Parsear el contenido del mensaje
                    msg = email.message_from_bytes(response[1])

                    # Obtener el asunto del correo
                    subject, encoding = decode_header(msg["Subject"])[0]
                    if isinstance(subject, bytes):
                        # Si el asunto está codificado en bytes, decodificarlo
                        subject = subject.decode(encoding if encoding else "utf-8")

                    # Obtener el remitente
                    from_ = msg.get("From")

                    # Verificar si el correo coincide con el filtro
                    if filter_email and filter_email not in from_:
                        continue

                    # Obtener el contenido del mensaje
                    if msg.is_multipart():
                        for part in msg.walk():
                            if part.get_content_type() == "text/plain":
                                body = part.get_payload(decode=True).decode("utf-8", errors="ignore")
                                break
                    else:
                        body = msg.get_payload(decode=True).decode("utf-8", errors="ignore")

                    # Limitar a 30 líneas
                    lines = body.splitlines()
                    body = "\n".join(lines[:30])

                    # Crear contenido con encabezado
                    content = f"De: {from_}\nAsunto: {subject}\n{'-' * 50}\n{body}"
                    content_list.append(content)

        # Guardar todos los mensajes en un solo archivo
        print(f"Guardando todos los mensajes en formato {format}...")
        save_messages_to_single_file(content_list, format, save_path)
        print(f"Mensajes guardados en: {save_path}/emails.{format}")

        # Cerrar la conexión y cerrar sesión
        imap.close()
        imap.logout()

    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    
    # read_emails(
    #     username=username,
    #     password=password,
    #     number_emails=10,
    #     format="word",  # Cambiar a "pdf" o "txt" según sea necesario
    #     filter_email="",  # Filtrar por remitente
    #     save_path="./documentos_generados/docx"  # Directorio donde guardar los archivos
    # )

    send_email(
        to_email="catiloji66@gmail.com",
        subject="Correo de prueba automatico",
        message="Este es un mensaje de prueba automatico generado por Auto-pythona de Pied Piper Inc, CEO: Alfonso almenara",
        cc_email="antonioalm45@gmail.com",
        bcc_email="alfonsoalm34@gmail.com"
    )
